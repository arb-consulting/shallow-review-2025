#!/usr/bin/env python3
"""
docx_links_to_csv.py

Extract all hyperlinks from a DOCX file and write them to a CSV file.

Usage:
  python docx_links_to_csv.py --docx <DOCX_FILE_PATH> [--out links.csv]

Requirements:
  pip install python-docx
"""

import argparse
import csv
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract_links_from_paragraph(paragraph: Paragraph) -> list[tuple[str, str]]:
    """
    Extract hyperlinks from a paragraph.
    Returns a list of tuples: (link_text, url)
    """
    links = []
    
    # Iterate through all hyperlink elements in the paragraph
    for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
        # Get the link text
        link_text = ''.join(node.text for node in hyperlink.xpath('.//w:t') if node.text)
        
        # Get the relationship ID
        r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        
        if r_id:
            # Get the actual URL from the document relationships
            try:
                url = paragraph.part.rels[r_id].target_ref
                links.append((link_text.strip(), url))
            except KeyError:
                # Handle case where relationship doesn't exist
                pass
    
    return links


def extract_links_from_table(table: Table) -> list[tuple[str, str]]:
    """Extract hyperlinks from a table."""
    links = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                links.extend(extract_links_from_paragraph(paragraph))
    return links


def extract_links_from_docx(docx_path: str) -> list[tuple[str, str]]:
    """
    Extract all hyperlinks from a DOCX file.
    Returns a list of tuples: (link_text, url)
    """
    doc = Document(docx_path)
    all_links = []
    
    # Extract links from paragraphs
    for paragraph in doc.paragraphs:
        all_links.extend(extract_links_from_paragraph(paragraph))
    
    # Extract links from tables
    for table in doc.tables:
        all_links.extend(extract_links_from_table(table))
    
    # Extract links from headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            all_links.extend(extract_links_from_paragraph(paragraph))
        for table in header.tables:
            all_links.extend(extract_links_from_table(table))
    
    # Extract links from footers
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            all_links.extend(extract_links_from_paragraph(paragraph))
        for table in footer.tables:
            all_links.extend(extract_links_from_table(table))
    
    return all_links


def main():
    parser = argparse.ArgumentParser(
        description="Extract all hyperlinks from a DOCX file to a CSV."
    )
    parser.add_argument(
        "--docx",
        default="./input/shallow-review-2025.docx",
        help="Path to the DOCX file to extract links from (default: ./input/shallow-review-2025.docx).",
    )
    parser.add_argument(
        "--out",
        default="results/shallow-review-2025-links.csv",
        help="Output CSV path (default: results/shallow-review-2025-links.csv).",
    )
    args = parser.parse_args()
    
    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        return 1
    
    print(f"Extracting links from {docx_path}...")
    links = extract_links_from_docx(str(docx_path))
    
    # Remove duplicates while preserving order
    seen = set()
    unique_links = []
    for link_text, url in links:
        if (link_text, url) not in seen:
            seen.add((link_text, url))
            unique_links.append((link_text, url))
    
    print(f"Found {len(unique_links)} unique links.")
    
    out_path = Path(args.out)
    with out_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["title", "url"])
        for link_text, url in unique_links:
            writer.writerow([link_text, url])
    
    print(f"Wrote {len(unique_links)} links to: {out_path.resolve()}")
    return 0


if __name__ == "__main__":
    exit(main())

